import docx
from drivingschool.models import Group, Weekday

def gen_report():
    doc = docx.Document()
    doc.add_heading('Отчет о расписании теоретических занятий', 0)

    for group in Group.objects.all():
        rows = group.scheduletheory_set.all()
        groupname = doc.add_paragraph(group.name)
        paragraph_format = groupname.paragraph_format
        table = doc.add_table(rows=8, cols=3)
        table.style = 'Table Grid'

        table.cell(0, 0).text = 'День недели'
        table.cell(0, 1).text = 'Время'
        table.cell(0, 2).text = 'Преподаватели'

        for i, weekday in enumerate(Weekday.choices, start=1):
            print(i, weekday)
            table.cell(i, 0).text = str(weekday[1])
            table.cell(i, 1).text = '---'
            table.cell(i, 2).text = '---'


        for i, schedule in enumerate(group.scheduletheory_set.all(), 1):
            print(i, schedule)
            table.cell(i, 0).text = schedule.get_weekday_display()
            table.cell(i, 1).text = schedule.get_position_display()
            table.cell(i, 2).text = schedule.substitute_tutor if schedule.substitute_tutor else '---'

    doc.save('table.docx')